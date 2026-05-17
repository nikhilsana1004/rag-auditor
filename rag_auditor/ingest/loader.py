"""
Document loader: ingests PDF, TXT, DOCX, and Markdown files
from a file path or directory.
"""

from pathlib import Path
from dataclasses import dataclass, field
from typing import List


@dataclass
class Document:
    content: str
    source: str
    metadata: dict = field(default_factory=dict)


SUPPORTED = {".txt", ".md", ".pdf", ".docx"}


def load_documents(path: Path) -> List[Document]:
    path = Path(path)
    files = _resolve_files(path)
    docs = []
    for f in files:
        docs.extend(_load_file(f))
    if not docs:
        raise ValueError(f"No supported documents found in: {path}")
    return docs


def _resolve_files(path: Path) -> List[Path]:
    if path.is_file():
        return [path]
    if path.is_dir():
        return sorted(f for f in path.rglob("*") if f.suffix.lower() in SUPPORTED)
    raise FileNotFoundError(f"Path not found: {path}")


def _load_file(path: Path) -> List[Document]:
    ext = path.suffix.lower()
    try:
        if ext in {".txt", ".md"}:
            return _load_text(path)
        elif ext == ".pdf":
            return _load_pdf(path)
        elif ext == ".docx":
            return _load_docx(path)
    except Exception as e:
        print(f"  [warn] Could not load {path.name}: {e}")
    return []


def _load_text(path: Path) -> List[Document]:
    content = path.read_text(encoding="utf-8", errors="ignore").strip()
    return [Document(content=content, source=str(path))]


def _load_pdf(path: Path) -> List[Document]:
    try:
        import pdfplumber
        docs = []
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    docs.append(Document(
                        content=text.strip(),
                        source=str(path),
                        metadata={"page": i + 1},
                    ))
        return docs
    except ImportError:
        # Fallback: pypdf
        import pypdf
        reader = pypdf.PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(content=text.strip(), source=str(path), metadata={"page": i + 1}))
        return docs


def _load_docx(path: Path) -> List[Document]:
    import docx
    doc = docx.Document(str(path))
    content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return [Document(content=content, source=str(path))]
